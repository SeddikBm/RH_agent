"""
Service de parsing de CVs — PDF et DOCX.
Extraction du texte brut uniquement.
Le parsing en sections est effectué par le LLM (voir section_extractor.py).
"""
import re
from pathlib import Path

from loguru import logger


# ── Extraction du texte brut ───────────────────────────────

def parse_pdf(file_path: str) -> str:
    """Extrait le texte d'un fichier PDF via PyMuPDF."""
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(file_path)
        text_parts = []

        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text("text")
            if text.strip():
                text_parts.append(text)

        doc.close()
        full_text = "\n".join(text_parts)
        logger.info(f"✅ PDF parsé : {file_path} ({len(full_text)} caractères)")
        return _clean_text(full_text)

    except Exception as e:
        logger.error(f"❌ Erreur parsing PDF {file_path}: {e}")
        raise


def parse_docx(file_path: str) -> str:
    """Extrait le texte d'un fichier DOCX via python-docx."""
    try:
        from docx import Document

        doc = Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

        # Extraire aussi les tableaux
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)

        full_text = "\n".join(paragraphs)
        logger.info(f"✅ DOCX parsé : {file_path} ({len(full_text)} caractères)")
        return _clean_text(full_text)

    except Exception as e:
        logger.error(f"❌ Erreur parsing DOCX {file_path}: {e}")
        raise


def parse_txt(file_path: str) -> str:
    """Extrait le texte d'un fichier TXT."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        return _clean_text(text)
    except Exception as e:
        logger.error(f"❌ Erreur parsing TXT {file_path}: {e}")
        raise


def parse_cv(file_path: str) -> str:
    """Point d'entrée unifié — détecte le format et parse."""
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext == ".pdf":
        return parse_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return parse_docx(file_path)
    elif ext == ".txt":
        return parse_txt(file_path)
    else:
        raise ValueError(f"Format de fichier non supporté : {ext}. Utilisez PDF, DOCX ou TXT.")


# ── Nettoyage du texte ─────────────────────────────────────

def _clean_text(text: str) -> str:
    """Nettoie le texte extrait : espaces multiples, caractères spéciaux."""
    # Supprimer les caractères de contrôle (sauf newlines et tabs)
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)

    # Normaliser les espaces multiples
    text = re.sub(r" {2,}", " ", text)

    # Normaliser les sauts de ligne multiples (max 2)
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()
